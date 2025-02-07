import codecs, random, time, itertools
from docx import Document
MOD = 256

# Функція KSA для подвійного потоку
def KSA(keys):
    key_lengths = [len(key) for key in keys]
    S = list(range(MOD))
    j = 0
    for i in range(MOD):
        key_sum = sum([ord(key[i % length]) for key, length in zip(keys, key_lengths)])
        j = (j + S[i] + key_sum) % MOD
        S[i], S[j] = S[j], S[i]
    return S

# Функція PRGA для генерації потоку
def PRGA(S):
    i = 0
    j = 0
    while True:
        i = (i + 1) % MOD
        j = (j + S[i]) % MOD
        S[i], S[j] = S[j], S[i]
        K = S[(S[i] + S[j]) % MOD]
        yield K

# Генерація потоку для двох ключів
def get_keystream(keys):
    S = KSA(keys)
    return PRGA(S)

# Логіка шифрування для англомовного тексту
def encrypt_logic_en(keys, text):
    keystream = get_keystream(keys)
    res = []
    for c in text:
        val = ("%02X" % (c ^ next(keystream)))
        res.append(val)
    encrypted_text = ''.join(res)
    if len(encrypted_text) % 2 != 0:
        encrypted_text = '0' + encrypted_text
    return encrypted_text

# Логіка шифрування для україномовного тексту
def encrypt_logic_ua(keys, text):
    keystream = get_keystream(keys)
    res = []
    for c in text:
        val = ("%03X" % (ord(c) ^ next(keystream)))
        res.append(val)
    encrypted_text = ''.join(res)
    return encrypted_text

# Шифрування для англійської мови
def encrypt_en(keys, text):
    text = [ord(c) for c in text]
    return encrypt_logic_en(keys, text)

# Шифрування для української мови
def encrypt_ua(keys, text):
    return encrypt_logic_ua(keys, text)

# Дешифрування для англійської мови
def decrypt_en(keys, text):
    text = codecs.decode(text, 'hex_codec')
    res = encrypt_logic_en(keys, text)
    return codecs.decode(res, 'hex_codec').decode('utf-8')

# Дешифрування для української мови
def decrypt_ua(keys, text):
    keystream = get_keystream(keys)
    res = []
    for i in range(0, len(text), 3):
        val = int(text[i:i + 3], 16)  # Отримуємо зашифроване значення
        decrypted_val = val ^ next(keystream)  # XOR з наступним значенням потоку
        res.append(chr(decrypted_val))
    decrypted_text = ''.join(res)
    return decrypted_text

# Генерація випадкових рядків для ключів з довжиною, яка становить 0.7 від довжини вхідного тексту
def generate_random_string(length, lang):
    alphabet = ''
    if lang == 'ua':
        alphabet = 'абвгґдеєжзиіїйклмнопрстуфхцчшщьюя'
    elif lang == 'en':
        alphabet = 'abcdefghijklmnopqrstuvwxyz'
    random_string = ''.join(random.choice(alphabet) for _ in range(length))
    return random_string

# Функція для визначення довжини ключа 0.7 від довжини тексту
def calculate_key_length(text_length):
    return int(text_length * 0.7)

# Шифрування для англійської мови
def en_text():
    doc = Document('input.docx')
    content = ""
    for paragraph in doc.paragraphs:
        content += paragraph.text
    
    print(f"\n------------------------------------------------------------")
    print("Input length = ", len(content))

    # Розрахунок довжини ключа
    key_length = calculate_key_length(len(content))
    print(f"Key length = {key_length} (0.7 of input length)")
    
    # Генерація ключів
    key1 = generate_random_string(key_length, 'en')
    key2 = generate_random_string(key_length, 'en')
    keys = [key1, key2]
    
    print("Keys:", keys)
    
    print("------------------------------------------------------------")
    print('input =', content)
    print("------------------------------------------------------------")
    
    start_time = time.time()
    encrypted_text = encrypt_en(keys, content)
    total_time = (time.time() - start_time)
    print('Total encode time =', total_time, 's')
    
    start_time = time.time()
    decrypted_text = decrypt_en(keys, encrypted_text)
    total_time = (time.time() - start_time)
    print('Total decode time =', total_time, 's')
    
    print('encrypt =', encrypted_text)
    print('decrypt =', decrypted_text)
    print(f"\n")
    
    # Збереження дешифрованого тексту у новому документі
    doc1 = Document()
    doc1.add_paragraph(decrypted_text)
    doc1.save("output.docx")

    # Зламування шифру за допомогою hacrypt
    hacrypt(keys, 'en')

# Шифрування для української мови
def ua_text():
    with open('input.txt', 'r', encoding='utf-8') as file:
        content = file.read()
        content = content[:-1]

    print(f"\n------------------------------------------------------------")
    print("Input length = ", len(content))

    # Розрахунок довжини ключа
    key_length = calculate_key_length(len(content))
    print(f"Key length = {key_length} (0.7 of input length)")
    
    # Генерація ключів
    key1 = generate_random_string(key_length, 'ua')
    key2 = generate_random_string(key_length, 'ua')
    keys = [key1, key2]
    
    print("Keys:", keys)
    
    print("------------------------------------------------------------")
    print('input =', content)
    print("------------------------------------------------------------")
    
    start_time = time.time()
    encrypted_text = encrypt_ua(keys, content)
    total_time = (time.time() - start_time)
    print('Total encode time =', total_time, 's')
    
    start_time = time.time()
    decrypted_text = decrypt_ua(keys, encrypted_text)
    total_time = (time.time() - start_time)
    print('Total decode time =', total_time, 's')
    
    print('encrypt =', encrypted_text)
    print('decrypt =', decrypted_text)
    print(f"\n")
    
    # Збереження дешифрованого тексту у новому документі
    doc1 = Document()
    doc1.add_paragraph(decrypted_text)
    doc1.save("output_ua.docx")

    # Зламування шифру за допомогою hacrypt
    hacrypt(keys, 'ua')

# Функція hacrypt для порівняння комбінацій
import itertools

def hacrypt(keys, lang):
    alphabet = ''
    if lang == 'ua':
        alphabet = 'абвгґдеєжзиіїйклмнопрстуфхцчшщьюя'
    elif lang == 'en':
        alphabet = 'abcdefghijklmnopqrstuvwxyz'
    start_time = time.time()
    length = 2  # Починаємо з двох символів в рядку
    while length <= 7:  # Збільшуємо довжину символів в рядку до 6 влкючно
        for combination in itertools.product(alphabet, repeat=length):
            word1 = ''.join(combination)
            for combination in itertools.product(alphabet, repeat=length):
                word2 = ''.join(combination)
                words = [word1, word2]  
                # print(words)
                if words == keys:
                    total_time = (time.time() - start_time)
                    print(f"SUCCESS with {words}. It has taken: {total_time} s")
                    
                    return
        length += 1  # Збільшоємо довжину рядка на 1




if __name__ == "__main__":
    # en_text()
    ua_text()

