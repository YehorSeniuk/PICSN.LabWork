import random
import time
import docx
from docx import Document

SYMBOLS = 'ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz0123456789 !?.,'


def get_random_key():
    while True:
        keyA = random.randint(2, len(SYMBOLS))
        keyB = random.randint(2, len(SYMBOLS))
        if compute_gcd(keyA, len(SYMBOLS)) == 1:
            return keyA, keyB


def compute_gcd(a, b):
    while a != 0:
        a, b = b % a, a
    return b


def find_modular_inverse(a, m):
    if compute_gcd(a, m) != 1:
        return None
    u, v, gcd = 1, 0, a
    prev_u, prev_v, prev_gcd = 0, 1, m
    while gcd != 0:
        quotient = prev_gcd // gcd
        prev_u, u = u, prev_u - quotient * u
        prev_v, v = v, prev_v - quotient * v
        prev_gcd, gcd = gcd, prev_gcd - quotient * gcd
    return prev_u % m if prev_gcd == 1 else None


def encrypt(keyA, keyB, text):
    result = ''
    for symbol in text:
        if symbol in SYMBOLS:
            sym_index = SYMBOLS.find(symbol)
            result += SYMBOLS[(sym_index * keyA + keyB) % len(SYMBOLS)]
        else:
            result += symbol
    return result


def decrypt(keyA, keyB, text):
    result = ''
    modular_inverse_of_keyA = find_modular_inverse(keyA, len(SYMBOLS))
    for symbol in text:
        if symbol in SYMBOLS:
            symbol_index = SYMBOLS.find(symbol)
            result += SYMBOLS[(symbol_index - keyB) * modular_inverse_of_keyA % len(SYMBOLS)]
        else:
            result += symbol
    return result


def haccrypt(text):
    for i in range(2, len(SYMBOLS)):
        for j in range(2, len(SYMBOLS)):
            if compute_gcd(i, len(SYMBOLS)) != 1:
                break
            encrypted_text = decrypt(i, j, text)


if __name__ == '__main__':
    keyA, keyB = get_random_key()
    print("keyA =", keyA, "keyB =", keyB)

    doc = Document('input.docx')
    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])

    print("------------------------------------------------------------")
    print('Input:', text)
    print('Input length:', len(text))
    print("------------------------------------------------------------\n")

    start_time = time.time()
    encrypted_text = encrypt(keyA, keyB, text)
    print('Total encode time =', time.time() - start_time, 's')

    start_time = time.time()
    decrypted_text = decrypt(keyA, keyB, encrypted_text)
    print('Total decode time =', time.time() - start_time, 's')

    print('Encrypted:', encrypted_text)
    print('Decrypted:', decrypted_text, "\n")

    doc1 = docx.Document()
    doc1.add_paragraph(decrypted_text)
    doc1.save("output.docx")

    print('-------------- Brute-force attack ----------------')
    start_time = time.time()
    haccrypt(encrypted_text)
    print('Total hack time =', time.time() - start_time, 's')

